from __future__ import annotations

import json
import logging
import re
from pathlib import Path

from app.shared.state_types import ParsedSection


log = logging.getLogger(__name__)

_DOCX_HEADING_PREFIX = "Heading"


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _parse_markdown(path: Path) -> list[ParsedSection]:
    text = _read_text(path)
    heading_re = re.compile(r"^(#{1,6})\s+(.+?)\s*$", re.MULTILINE)
    matches = list(heading_re.finditer(text))

    if not matches:
        return [{
            "file_name": path.name,
            "section_heading": None,
            "page": None,
            "content_type": "text",
            "content": text.strip(),
            "raw_image_ref": None,
        }]

    sections: list[ParsedSection] = []

    preamble = text[: matches[0].start()].strip()
    if preamble:
        sections.append({
            "file_name": path.name,
            "section_heading": None,
            "page": None,
            "content_type": "text",
            "content": preamble,
            "raw_image_ref": None,
        })

    for i, m in enumerate(matches):
        body_start = m.end()
        body_end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[body_start:body_end].strip()
        sections.append({
            "file_name": path.name,
            "section_heading": m.group(2).strip(),
            "page": None,
            "content_type": "text",
            "content": body,
            "raw_image_ref": None,
        })

    return sections


def _parse_plain_text(path: Path) -> list[ParsedSection]:
    return [{
        "file_name": path.name,
        "section_heading": None,
        "page": None,
        "content_type": "text",
        "content": _read_text(path).strip(),
        "raw_image_ref": None,
    }]


def _parse_csv(path: Path) -> list[ParsedSection]:
    return [{
        "file_name": path.name,
        "section_heading": None,
        "page": None,
        "content_type": "table",
        "content": _read_text(path).strip(),
        "raw_image_ref": None,
    }]


def _parse_json(path: Path) -> list[ParsedSection]:
    text = _read_text(path)
    try:
        pretty = json.dumps(json.loads(text), indent=2, ensure_ascii=False)
    except json.JSONDecodeError:
        pretty = text
    return [{
        "file_name": path.name,
        "section_heading": None,
        "page": None,
        "content_type": "text",
        "content": pretty.strip(),
        "raw_image_ref": None,
    }]


def _parse_pdf(path: Path) -> list[ParsedSection]:
    import fitz  # PyMuPDF

    sections: list[ParsedSection] = []
    with fitz.open(path) as doc:
        for page_num, page in enumerate(doc, start=1):
            text = (page.get_text() or "").strip()
            if not text:
                continue
            sections.append({
                "file_name": path.name,
                "section_heading": None,
                "page": page_num,
                "content_type": "text",
                "content": text,
                "raw_image_ref": None,
            })
    if not sections:
        return [{
            "file_name": path.name,
            "section_heading": None,
            "page": None,
            "content_type": "text",
            "content": "",
            "raw_image_ref": None,
        }]
    return sections


def _parse_docx(path: Path) -> list[ParsedSection]:
    from docx import Document

    document = Document(path)
    sections: list[ParsedSection] = []
    current_heading: str | None = None
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            sections.append({
                "file_name": path.name,
                "section_heading": current_heading,
                "page": None,
                "content_type": "text",
                "content": "\n\n".join(buffer).strip(),
                "raw_image_ref": None,
            })
            buffer.clear()

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "") if para.style else ""
        if style.startswith(_DOCX_HEADING_PREFIX):
            flush()
            current_heading = text
        else:
            buffer.append(text)
    flush()

    for i, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append("\t".join((cell.text or "").strip() for cell in row.cells))
        body = "\n".join(rows).strip()
        if body:
            sections.append({
                "file_name": path.name,
                "section_heading": f"Table {i}",
                "page": None,
                "content_type": "table",
                "content": body,
                "raw_image_ref": None,
            })

    return sections or [{
        "file_name": path.name,
        "section_heading": None,
        "page": None,
        "content_type": "text",
        "content": "",
        "raw_image_ref": None,
    }]


_PARSERS = {
    ".md": _parse_markdown,
    ".markdown": _parse_markdown,
    ".txt": _parse_plain_text,
    ".csv": _parse_csv,
    ".json": _parse_json,
    ".pdf": _parse_pdf,
    ".docx": _parse_docx,
}


def parse_file(path: Path) -> list[ParsedSection]:
    """Parse a supported file and return normalized sections.

    Supported: .md/.markdown, .txt, .csv, .json, .pdf, .docx.
    Unknown extensions fall back to plain-text best-effort.
    """
    suffix = path.suffix.lower()
    parser = _PARSERS.get(suffix, _parse_plain_text)
    fallback = parser is _parse_plain_text and suffix not in _PARSERS
    log.info("[parser] parse_file path=%s ext=%s parser=%s fallback=%s", path.name, suffix, parser.__name__, fallback)
    sections = parser(path)
    log.info("[parser] parse_file DONE path=%s sections=%d", path.name, len(sections))
    return sections
