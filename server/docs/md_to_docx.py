"""
Convert sprint_agent.md to sprint_agent.docx with proper heading styles,
tables, code blocks, and bullet lists.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_code_block(doc: Document, code: str):
    """Add a monospace shaded paragraph for code/pre blocks."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    # grey background via shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    run = para.add_run(code)
    run.font.name  = "Courier New"
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)


def inline_format(para, text: str):
    """
    Parse inline markdown in `text` and add formatted runs to `para`.
    Supports **bold**, `code`, and plain text.
    """
    pattern = re.compile(r"(\*\*(.+?)\*\*|`([^`]+)`)")
    pos = 0
    for m in pattern.finditer(text):
        # plain text before match
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        if m.group(2):           # **bold**
            run = para.add_run(m.group(2))
            run.bold = True
        elif m.group(3):         # `code`
            run = para.add_run(m.group(3))
            run.font.name  = "Courier New"
            run.font.size  = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])


# ── main converter ──────────────────────────────────────────────────────────

def convert(md_path: Path, docx_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code_block   = False
    code_lines: list[str] = []
    in_table        = False
    table_rows: list[list[str]] = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        col_count = max(len(r) for r in table_rows)
        tbl = doc.add_table(rows=len(table_rows), cols=col_count)
        tbl.style = "Table Grid"
        for r_idx, row in enumerate(table_rows):
            for c_idx, cell_text in enumerate(row):
                cell = tbl.cell(r_idx, c_idx)
                cell.text = ""
                p = cell.paragraphs[0]
                inline_format(p, cell_text.strip())
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
                    set_cell_bg(cell, "1F2937")
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif r_idx % 2 == 0:
                    set_cell_bg(cell, "F9FAFB")
        doc.add_paragraph()
        in_table = False
        table_rows.clear()

    while i < len(lines):
        line = lines[i]

        # ── code fence ──────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lines = []
            else:
                add_code_block(doc, "\n".join(code_lines))
                in_code_block = False
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── table ───────────────────────────────────────────────
        if line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # skip separator row (|---|---|)
            if not all(re.match(r"^[-: ]+$", c) for c in cells):
                table_rows.append(cells)
            in_table = True
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ── horizontal rule ─────────────────────────────────────
        if line.strip() in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── headings ────────────────────────────────────────────
        m = re.match(r"^(#{1,6})\s+(.*)", line)
        if m:
            level  = len(m.group(1))
            text   = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", m.group(2))  # strip links
            heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3",
                           4: "Heading 4", 5: "Heading 5", 6: "Heading 6"}
            p = doc.add_heading(text, level=min(level, 6))
            i += 1
            continue

        # ── bullet list ─────────────────────────────────────────
        m = re.match(r"^(\s*)[-*+]\s+(.*)", line)
        if m:
            indent = len(m.group(1)) // 2
            text   = m.group(2)
            style  = "List Bullet 2" if indent > 0 else "List Bullet"
            p = doc.add_paragraph(style=style)
            inline_format(p, text)
            i += 1
            continue

        # ── numbered list ────────────────────────────────────────
        m = re.match(r"^\s*\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            inline_format(p, m.group(1))
            i += 1
            continue

        # ── blockquote ───────────────────────────────────────────
        m = re.match(r"^>\s+(.*)", line)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            run = p.add_run(m.group(1))
            run.italic = True
            run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            i += 1
            continue

        # ── empty line ───────────────────────────────────────────
        if line.strip() == "":
            i += 1
            continue

        # ── normal paragraph ─────────────────────────────────────
        p = doc.add_paragraph()
        inline_format(p, re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", line))
        i += 1

    if in_table:
        flush_table()

    doc.save(str(docx_path))
    print(f"✅  Saved: {docx_path}")


if __name__ == "__main__":
    base = Path(__file__).parent
    convert(base / "sprint_agent.md", base / "sprint_agent.docx")

